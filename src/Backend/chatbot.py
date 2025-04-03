import os
import faiss
import torch
from fastapi import APIRouter, UploadFile, File, Form
from sentence_transformers import SentenceTransformer
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore

from langchain.schema import Document
from llama_cpp import Llama
import pdfplumber
from docx import Document as DocxDocument
from typing import Dict
from langchain_huggingface import HuggingFaceEmbeddings

router = APIRouter()

# Ensure temp directory exists
os.makedirs("temp", exist_ok=True)

# Load embedding model
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Determine embedding dimension dynamically
test_embedding = embedding_model.embed_query("test")
dimension = len(test_embedding)  # Get embedding size from a dummy query

# Create FAISS index
faiss_index = faiss.IndexFlatL2(dimension)

# Initialize FAISS vector store
vector_store = FAISS(
    embedding_function=embedding_model,
    index=faiss_index,
    docstore=InMemoryDocstore(),
    index_to_docstore_id={}
)

# Load Llama model
llama = Llama(model_path="models/llama-7b.Q4_K_M.gguf", n_ctx=2048)


# Function to process and embed documents
def process_document(file_path):
    ext = file_path.split('.')[-1].lower()
    text = ""

    if ext == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif ext == "pdf":
        with pdfplumber.open(file_path) as pdf:
            extracted_text = [page.extract_text() for page in pdf.pages if page.extract_text()]
            text = "\n".join(extracted_text) if extracted_text else ""
    elif ext == "docx":
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs]) if doc.paragraphs else ""

    if text.strip():
        document = Document(page_content=text)
        vector_store.add_documents([document])
        return {"message": "Document processed successfully!"}
    
    return {"error": "Unsupported or empty file format!"}


# Endpoint for document upload and processing
@router.post("/upload/")
async def upload_document(file: UploadFile = File(...)):
    file_path = f"temp/{file.filename}"
    with open(file_path, "wb") as f:
        f.write(await file.read())

    return process_document(file_path)


# RAG Query Endpoint
@router.post("/query/")
async def rag_query(user_input: str = Form(...)) -> Dict[str, str]:
    retrieved_docs = vector_store.similarity_search(user_input, k=3)
    context = "\n".join([doc.page_content for doc in retrieved_docs])

    try:
        response = llama.create_completion(
            f"Context: {context}\n\nQuestion: {user_input}\nAnswer:",
            max_tokens=200
        )
        answer = response["choices"][0]["text"].strip()
    except Exception as e:
        return {"error": f"Failed to generate response: {str(e)}"}

    return {"response": answer}
