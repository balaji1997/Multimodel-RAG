import os
import faiss
import streamlit as st
import torch
from sentence_transformers import SentenceTransformer
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.docstore import InMemoryDocstore
from langchain.schema import Document
from llama_cpp import Llama
from pdfplumber import open as pdf_open
from docx import Document as DocxDocument

# Ensure temp directory exists
os.makedirs("temp", exist_ok=True)

# Load local embedding model
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Create FAISS index
dimension = embedding_model.client.get_sentence_embedding_dimension()
faiss_index = faiss.IndexFlatL2(dimension)

# Initialize FAISS vector store
vector_store = FAISS(
    embedding_function=embedding_model,
    index=faiss_index,
    docstore=InMemoryDocstore(),
    index_to_docstore_id={}
)

# Load local Llama model
llama = Llama(model_path="models/llama-7b.Q4_K_M.gguf", n_ctx=2048)

print(llama)

# Function to process and embed documents
def process_document(file_path):
    ext = file_path.split('.')[-1]
    if ext == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif ext == "pdf":
        with pdf_open(file_path) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif ext == "docx":
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        return "Unsupported file format!"
    
    document = Document(page_content=text)
    vector_store.add_documents([document])
    return "Document processed successfully!"

# RAG Query Function
def rag_query(user_input):
    retrieved_docs = vector_store.similarity_search(user_input, k=3)
    context = "\n".join([doc.page_content for doc in retrieved_docs])
    
    response = llama(
        f"Context: {context}\n\nQuestion: {user_input}\nAnswer:",
        max_tokens=200
    )
    return response["choices"][0]["text"].strip()

# Streamlit UI
def main():
    st.title("Open-Source RAG Knowledge Assistant")
    st.sidebar.header("Upload Documents")
    uploaded_file = st.sidebar.file_uploader("Upload a file", type=["txt", "pdf", "docx"])

    if uploaded_file:
        file_path = f"temp/{uploaded_file.name}"
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        process_document(file_path)
        st.sidebar.success("File processed successfully!")

    user_input = st.text_input("Ask a question:")
    if st.button("Submit"):
        if user_input.strip():
            response = rag_query(user_input)
            st.write("Response:", response)
        else:
            st.warning("Please enter a question.")

if __name__ == "__main__":
    main()